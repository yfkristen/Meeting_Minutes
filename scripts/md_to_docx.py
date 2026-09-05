#!/usr/bin/env python3
"""將 minutes/ 的會議紀錄 Markdown 轉成 Word (.docx)。

用法:
    python3 scripts/md_to_docx.py minutes/2026-08-20_雙週會.md
    python3 scripts/md_to_docx.py minutes/2026-08-20_雙週會.md --email
    python3 scripts/md_to_docx.py minutes/2026-08-20_雙週會.md -o output/word/自訂.docx

輸出預設為 output/word/<原檔名>_會議紀錄.docx

Markdown 依 skill「ai-pmo-meeting-minutes」的格式撰寫：
    第一行為會議正式名稱，其下為 壹貳參 → 一二三 → 1. 2. 3. 三層編號。
紀錄本體不含郵件套語；需要連同郵件正文一起輸出時加 --email。
"""

import argparse
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# ---- 版面設定（要換字型/字級改這裡即可）----
CN_FONT = "標楷體"           # 中文字型
EN_FONT = "Times New Roman"  # 英數字型
BODY_SIZE = Pt(12)
TITLE_SIZE = Pt(15)
LINE_SPACING = 1.5

LEVEL1 = re.compile(r"^[壹貳參肆伍陸柒捌玖拾]+、")   # 壹、貳、參
LEVEL2 = re.compile(r"^[一二三四五六七八九十]+、")    # 一、二、三
LEVEL3 = re.compile(r"^(?:\d+\.|\(\d+\))\s*")       # 1. 2. 3. 或 (1) (2)
BULLET = re.compile(r"^[*\-\u2022]\s+")                # * - 項目符號
MAIL_HEAD = ("各位主管好", "檢送")
MAIL_TAIL = ("若有任何問題", "Best,")
WEEKDAYS = "一二三四五六日"

EMAIL_OPENING = "各位主管好：\n\n檢送{date}「{meeting}」會議紀錄，敬請參閱。"
EMAIL_CLOSING = "若有任何問題，再請聯繫我，謝謝！\n\nBest,"


def set_run_font(run, *, bold=False, size=BODY_SIZE):
    run.bold = bold
    run.font.size = size
    run.font.name = EN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)


def add_para(doc, text, *, bold=False, size=BODY_SIZE, center=False,
             indent_cm=0.0, hanging_cm=0.0, space_before=0, space_after=6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        pf.left_indent = Cm(indent_cm + hanging_cm)
        if hanging_cm:
            pf.first_line_indent = Cm(-hanging_cm)
    set_run_font(p.add_run(text), bold=bold, size=size)
    return p


def meeting_date(md_path: Path) -> str:
    """由檔名 YYYY-MM-DD_名稱 推出「2026/8/6(四)」格式。"""
    m = re.match(r"(\d{4})-(\d{2})-(\d{2})", md_path.stem)
    if not m:
        return ""
    y, mo, d = (int(g) for g in m.groups())
    return f"{y}/{mo}/{d}({WEEKDAYS[date(y, mo, d).weekday()]})"


def clean_lines(raw_lines: list[str]) -> list[str]:
    """去掉範本註解、水平線，以及已寫在 md 裡的郵件套語。"""
    out = []
    for raw in raw_lines:
        line = raw.strip()
        if line.startswith(">") or line == "---":
            continue
        if line.startswith(MAIL_HEAD) or line.startswith(MAIL_TAIL):
            continue
        out.append(line)
    while out and not out[0]:
        out.pop(0)
    while out and not out[-1]:
        out.pop()
    return out


def build(md_path: Path, out_path: Path, *, with_email: bool = False) -> Path:
    lines = clean_lines(md_path.read_text(encoding="utf-8").splitlines())

    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.54)
    section.left_margin = section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = EN_FONT
    normal.font.size = BODY_SIZE
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    # 第一行若不是編號段落，視為會議正式名稱
    title = ""
    if lines and not (LEVEL1.match(lines[0]) or LEVEL2.match(lines[0])
                      or LEVEL3.match(lines[0])):
        title = lines.pop(0)

    if with_email:
        opening = EMAIL_OPENING.format(
            date=meeting_date(md_path), meeting=title or md_path.stem)
        for para in opening.split("\n"):
            add_para(doc, para)
        add_para(doc, "", space_after=0)

    if title:
        add_para(doc, title, bold=True, size=TITLE_SIZE, center=True,
                 space_after=12)

    for line in lines:
        if not line:                           # 空行 → 段落間距
            add_para(doc, "", space_after=0)
        elif LEVEL1.match(line):               # 壹、貳、參
            add_para(doc, line, bold=True, hanging_cm=1.2, space_before=6)
        elif LEVEL2.match(line):               # 一、二、三
            add_para(doc, line, indent_cm=0.85, hanging_cm=1.2)
        elif LEVEL3.match(line):               # 1. 2. 3. 或 (1) (2)
            add_para(doc, line, indent_cm=1.7, hanging_cm=0.9)
        elif BULLET.match(line):               # 項目符號 → ‧
            add_para(doc, "‧" + BULLET.sub("", line),
                     indent_cm=2.6, hanging_cm=0.5, space_after=4)
        else:
            add_para(doc, line)

    if with_email:
        add_para(doc, "", space_after=0)
        for para in EMAIL_CLOSING.split("\n"):
            add_para(doc, para)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def main() -> int:
    ap = argparse.ArgumentParser(description="會議紀錄 Markdown → Word")
    ap.add_argument("markdown", help="minutes/ 底下的會議紀錄 .md")
    ap.add_argument("-o", "--output", help="輸出 .docx 路徑")
    ap.add_argument("--email", action="store_true",
                    help="連同郵件套語一起輸出（預設只輸出紀錄本體）")
    args = ap.parse_args()

    md_path = Path(args.markdown)
    if not md_path.is_file():
        print(f"找不到檔案：{md_path}", file=sys.stderr)
        return 1

    out_path = Path(args.output) if args.output else \
        Path("output/word") / f"{md_path.stem}_會議紀錄.docx"

    print(f"已產出：{build(md_path, out_path, with_email=args.email)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
